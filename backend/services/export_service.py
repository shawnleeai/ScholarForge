"""
文档导出服务
支持Word、PDF、Markdown格式导出
"""

from datetime import datetime
from typing import Optional, List, Dict, Any
from io import BytesIO
import re


class ExportService:
    """文档导出服务"""

    def __init__(self):
        self.supported_formats = ['markdown', 'docx', 'pdf']

    async def export_literature_review(
        self,
        review_data: Dict[str, Any],
        format: str = 'markdown',
    ) -> bytes:
        """
        导出文献综述

        Args:
            review_data: 综述数据
            format: 导出格式 (markdown, docx, pdf)

        Returns:
            文件内容 bytes
        """
        if format == 'markdown':
            return self._export_to_markdown(review_data)
        elif format == 'docx':
            return await self._export_to_docx(review_data)
        elif format == 'pdf':
            return await self._export_to_pdf(review_data)
        else:
            raise ValueError(f"不支持的格式: {format}")

    def _export_to_markdown(self, data: Dict[str, Any]) -> bytes:
        """导出为Markdown"""
        lines = []

        # 标题
        lines.append(f"# {data.get('title', '文献综述')}")
        lines.append('')

        # 摘要
        if data.get('abstract'):
            lines.append('## 摘要')
            lines.append('')
            lines.append(data['abstract'])
            lines.append('')

        # 关键词
        if data.get('keywords'):
            lines.append(f"**关键词：** {', '.join(data['keywords'])}")
            lines.append('')

        # 章节
        for section in data.get('sections', []):
            lines.append(f"## {section.get('title', '')}")
            lines.append('')
            lines.append(section.get('content', ''))
            lines.append('')

            # 子章节
            for sub in section.get('subsections', []):
                lines.append(f"### {sub.get('title', '')}")
                lines.append('')
                lines.append(sub.get('content', ''))
                lines.append('')

        # 研究空白
        if data.get('research_gaps'):
            lines.append('## 研究空白')
            lines.append('')
            for gap in data['research_gaps']:
                lines.append(f"- {gap}")
            lines.append('')

        # 未来方向
        if data.get('future_directions'):
            lines.append('## 未来研究方向')
            lines.append('')
            for direction in data['future_directions']:
                lines.append(f"- {direction}")
            lines.append('')

        # 参考文献
        if data.get('references'):
            lines.append('## 参考文献')
            lines.append('')
            for i, ref in enumerate(data['references'], 1):
                authors = ', '.join(ref.get('authors', [])[:3])
                if len(ref.get('authors', [])) > 3:
                    authors += ' et al.'
                lines.append(f"[{i}] {authors}. {ref.get('title', '')}. {ref.get('journal', '')} {ref.get('year', '')}.")
            lines.append('')

        # 元信息
        lines.append('---')
        lines.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
        lines.append(f"字数统计：{data.get('word_count', 0)} 字")

        content = '\n'.join(lines)
        return content.encode('utf-8')

    async def _export_to_docx(self, data: Dict[str, Any]) -> bytes:
        """导出为Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(12)

            # 标题
            title = doc.add_heading(data.get('title', '文献综述'), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 摘要
            if data.get('abstract'):
                doc.add_heading('摘要', 1)
                doc.add_paragraph(data['abstract'])

            # 章节
            for section in data.get('sections', []):
                doc.add_heading(section.get('title', ''), 1)
                doc.add_paragraph(section.get('content', ''))

                for sub in section.get('subsections', []):
                    doc.add_heading(sub.get('title', ''), 2)
                    doc.add_paragraph(sub.get('content', ''))

            # 参考文献
            if data.get('references'):
                doc.add_heading('参考文献', 1)
                for ref in data['references']:
                    authors = ', '.join(ref.get('authors', [])[:3])
                    if len(ref.get('authors', [])) > 3:
                        authors += ' et al.'
                    doc.add_paragraph(
                        f"{authors}. {ref.get('title', '')}. {ref.get('journal', '')} {ref.get('year', '')}.",
                        style='List Number'
                    )

            # 保存到内存
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        except ImportError:
            # 如果没有python-docx，返回Markdown
            return self._export_to_markdown(data)

    async def _export_to_pdf(self, data: Dict[str, Any]) -> bytes:
        """导出为PDF"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # 注册中文字体
            try:
                pdfmetrics.registerFont(TTFont('SimSun', 'simsun.ttc'))
            except:
                pass

            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)

            styles = getSampleStyleSheet()
            story = []

            # 标题
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1,  # 居中
            )
            story.append(Paragraph(data.get('title', '文献综述'), title_style))
            story.append(Spacer(1, 0.2 * inch))

            # 章节
            for section in data.get('sections', []):
                story.append(Paragraph(section.get('title', ''), styles['Heading2']))
                story.append(Spacer(1, 0.1 * inch))
                story.append(Paragraph(section.get('content', ''), styles['Normal']))
                story.append(Spacer(1, 0.2 * inch))

            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()

        except ImportError:
            # 如果没有reportlab，返回Markdown
            return self._export_to_markdown(data)

    async def export_paper(
        self,
        paper_data: Dict[str, Any],
        format: str = 'docx',
    ) -> bytes:
        """导出论文"""
        # TODO: 实现论文导出
        return b''


# 全局导出服务实例
export_service = ExportService()
